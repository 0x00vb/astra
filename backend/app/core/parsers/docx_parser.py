"""DOCX parser using python-docx."""
import io
import logging
from typing import Optional
from docx import Document as DocxDocument
from app.core.parsers.base import BaseParser, ParsedDocument

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """DOCX document parser."""

    def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        """
        Parse DOCX file.

        Args:
            file_content: Raw DOCX bytes
            filename: Original filename

        Returns:
            ParsedDocument with text and metadata
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
            }

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        paragraphs.append(" | ".join(row_texts))

            full_text = "\n\n".join(paragraphs)

            if not full_text.strip():
                raise ValueError("No text content extracted from DOCX")

            # Estimate pages (rough approximation: ~500 words per page)
            word_count = len(full_text.split())
            estimated_pages = max(1, word_count // 500)
            metadata["total_pages"] = estimated_pages

            return ParsedDocument(
                text=full_text,
                metadata=metadata,
                pages=None,  # DOCX doesn't have explicit page boundaries
            )

        except Exception as e:
            logger.error(f"Error parsing DOCX {filename}: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

